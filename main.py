from fastapi import FastAPI, HTTPException, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel
from deep_translator import GoogleTranslator
from docx import Document
import uvicorn
import io
import os
import tempfile

app = FastAPI(title="Language Translation Tool", version="5.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Hardcoded, definitive language dictionary ---
# Maps display name -> Google Translate ISO code
SUPPORTED_LANGUAGES: dict[str, str] = {
    "Afrikaans": "af",
    "Albanian": "sq",
    "Arabic": "ar",
    "Armenian": "hy",
    "Azerbaijani": "az",
    "Basque": "eu",
    "Belarusian": "be",
    "Bengali": "bn",
    "Bosnian": "bs",
    "Bulgarian": "bg",
    "Catalan": "ca",
    "Chinese (Simplified)": "zh-CN",
    "Chinese (Traditional)": "zh-TW",
    "Croatian": "hr",
    "Czech": "cs",
    "Danish": "da",
    "Dutch": "nl",
    "English": "en",
    "Esperanto": "eo",
    "Estonian": "et",
    "Finnish": "fi",
    "French": "fr",
    "Galician": "gl",
    "Georgian": "ka",
    "German": "de",
    "Greek": "el",
    "Gujarati": "gu",
    "Haitian Creole": "ht",
    "Hebrew": "iw",
    "Hindi": "hi",
    "Hungarian": "hu",
    "Icelandic": "is",
    "Indonesian": "id",
    "Irish": "ga",
    "Italian": "it",
    "Japanese": "ja",
    "Kannada": "kn",
    "Kazakh": "kk",
    "Korean": "ko",
    "Latin": "la",
    "Latvian": "lv",
    "Lithuanian": "lt",
    "Macedonian": "mk",
    "Malay": "ms",
    "Maltese": "mt",
    "Maori": "mi",
    "Marathi": "mr",
    "Mongolian": "mn",
    "Nepali": "ne",
    "Norwegian": "no",
    "Persian": "fa",
    "Polish": "pl",
    "Portuguese": "pt",
    "Punjabi": "pa",
    "Romanian": "ro",
    "Russian": "ru",
    "Serbian": "sr",
    "Slovak": "sk",
    "Slovenian": "sl",
    "Spanish": "es",
    "Swahili": "sw",
    "Swedish": "sv",
    "Tamil": "ta",
    "Telugu": "te",
    "Thai": "th",
    "Turkish": "tr",
    "Ukrainian": "uk",
    "Urdu": "ur",
    "Vietnamese": "vi",
    "Welsh": "cy",
    "Xhosa": "xh",
    "Yiddish": "yi",
    "Yoruba": "yo",
    "Zulu": "zu",
}


class TranslationRequest(BaseModel):
    text: str
    source_language: str   # ISO code, e.g. "en"
    target_language: str   # ISO code, e.g. "fr"


class TranslationResponse(BaseModel):
    translated_text: str
    source_language: str        # the code that was sent (may be "auto")
    detected_language: str      # the actual detected code (same as source unless auto)
    target_language: str


@app.get("/languages")
def get_languages():
    """Returns the hardcoded map of display names to ISO codes."""
    return SUPPORTED_LANGUAGES


@app.post("/translate", response_model=TranslationResponse)
def translate(request: TranslationRequest):
    valid_codes = set(SUPPORTED_LANGUAGES.values()) | {"auto"}

    # Validate source — "auto" is allowed for auto-detection
    if request.source_language not in valid_codes:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported source language code: '{request.source_language}'"
        )
    if request.target_language not in valid_codes or request.target_language == "auto":
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported target language code: '{request.target_language}'"
        )
    if not request.text.strip():
        raise HTTPException(status_code=400, detail="Input text cannot be empty.")
    if len(request.text) > 5000:
        raise HTTPException(status_code=400, detail="Input text exceeds 5000 character limit.")

    try:
        translator = GoogleTranslator(
            source=request.source_language,
            target=request.target_language
        )
        result: str = translator.translate(request.text.strip())
        if not result:
            raise ValueError("Translator returned an empty result.")

        # Resolve the detected language code
        detected = request.source_language
        if request.source_language == "auto":
            try:
                detected = GoogleTranslator(source="auto", target=request.target_language).detect(request.text.strip()) or "auto"
            except Exception:
                detected = "auto"  # fallback gracefully

        return TranslationResponse(
            translated_text=result,
            source_language=request.source_language,
            detected_language=detected,
            target_language=request.target_language,
        )
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Translation failed: {str(e)}"
        )


@app.post("/translate-document")
async def translate_document(
    file: UploadFile = File(...),
    target_language: str = Form(...),
    source_language: str = Form("auto")
):
    """
    Translate a .txt or .docx file.
    Returns the translated file for download.
    """
    if not target_language:
        raise HTTPException(status_code=400, detail="target_language is required")

    valid_codes = set(SUPPORTED_LANGUAGES.values()) | {"auto"}
    if source_language not in valid_codes:
        raise HTTPException(status_code=400, detail=f"Unsupported source language: {source_language}")
    if target_language not in valid_codes or target_language == "auto":
        raise HTTPException(status_code=400, detail=f"Unsupported target language: {target_language}")

    # Check file type
    if not file.filename:
        raise HTTPException(status_code=400, detail="No filename provided")

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.docx']:
        raise HTTPException(status_code=400, detail="Only .txt and .docx files are supported")

    try:
        contents = await file.read()

        if ext == '.txt':
            translated_content = translate_txt(contents, source_language, target_language)
            # Return as .txt
            output = io.BytesIO(translated_content.encode('utf-8'))
            filename = os.path.splitext(file.filename)[0] + f"_{target_language}.txt"
            return StreamingResponse(
                iter([output.getvalue()]),
                media_type="text/plain",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )

        elif ext == '.docx':
            translated_bytes = translate_docx(contents, source_language, target_language)
            filename = os.path.splitext(file.filename)[0] + f"_{target_language}.docx"
            return StreamingResponse(
                iter([translated_bytes]),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Document translation failed: {str(e)}")


def translate_txt(contents: bytes, source_lang: str, target_lang: str) -> str:
    """Translate plain text file."""
    text = contents.decode('utf-8', errors='replace').strip()
    if not text:
        return ""

    # Translate in chunks (max 5000 chars per request to avoid limits)
    chunks = []
    current_chunk = ""
    for line in text.split('\n'):
        if len(current_chunk) + len(line) + 1 > 4500:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = line
        else:
            current_chunk += ('\n' if current_chunk else '') + line

    if current_chunk:
        chunks.append(current_chunk)

    translator = GoogleTranslator(source=source_lang, target=target_lang)
    translated_chunks = [translator.translate(chunk) for chunk in chunks]
    return '\n'.join(translated_chunks)


def translate_docx(contents: bytes, source_lang: str, target_lang: str) -> bytes:
    """Translate a Word document, preserving formatting."""
    doc = Document(io.BytesIO(contents))
    translator = GoogleTranslator(source=source_lang, target=target_lang)

    # Translate all paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            para.text = translator.translate(para.text)

    # Translate tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        para.text = translator.translate(para.text)

    # Write to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


# Serve the frontend
app.mount("/static", StaticFiles(directory="."), name="static")

@app.get("/")
def serve_frontend():
    return FileResponse("index.html")


if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
